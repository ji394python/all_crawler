from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import urllib
import tqdm
import time
import pandas as pd
from docx import Document
from docx.shared import Inches

options = webdriver.ChromeOptions()
options.add_argument('--ignore-certificate-errors')
options.add_argument('--ignore-ssl-errors')
head_url = 'https://www.tandfonline.com'
if __name__ == '__main__':
    total_df = pd.DataFrame(columns=["類別","作者","標題","出版日","摘要","關鍵字"])
    driver = webdriver.Chrome(executable_path='C:\\Users\\chiaming\\Desktop\\vs\\武漢肺炎專區\\chromedriver.exe',chrome_options=options)
    driver.get('https://www.tandfonline.com/toc/uasa20/114/525?nav=tocList')
    html = BeautifulSoup(WebDriverWait(driver,10).until(lambda b:b.page_source),'lxml')
    url = [ head_url+i['href'] for i in html.select('.tocContent .articleEntry .art_title > a') ] 
    
    Title = []
    Author = []
    Date = []
    Keyword = []
    Kind = []
    Abstract = []
    
    for u in url:
        print(url.index(u))
        driver.get(u)
        time.sleep(1)
        second = BeautifulSoup(WebDriverWait(driver,10).until(lambda b:b.page_source),'lxml')
        main = second.select('.abstractInFull p')
        if len(main) == 0:
            continue
        else:
            author = second.select('.hlFld-ContribAuthor')[0].text.replace('show all','').strip()
            print('author:',author)
            try:
                keyword = second.select('.hlFld-KeywordText')[0].text.replace('KEYWORDS: ','').replace('\xa0','').strip()
            except:
                keyword = ""
            print('keyword:',keyword)
            kind = second.select('.toc-heading h3')[0].text.strip()
            print('kind:',kind)
            title = second.select('.hlFld-title')[0].text.strip()
            print('title:',title)
            print('【Date】')
            try:
                date = second.select('.itemPageRangeHistory')[0].text
            except:
                second = BeautifulSoup(WebDriverWait(driver,20).until(lambda b:b.page_source),'lxml')
                date = second.select('.itemPageRangeHistory')[0].text
            date = date[date.find('Published online: '):].replace('Published online: ','').strip()
            print("date:",date)

        Title.append(title)
        Author.append(author)
        Kind.append(kind)
        Date.append(date)
        Abstract.append(main)
        Keyword.append(keyword)
    new = {"類別":Kind,"作者":Author,"標題":Title,"出版日":Date,"摘要":Abstract,"關鍵字":Keyword}
    new = pd.DataFrame(new)
    print(new.head(3))
    total_df = pd.concat([total_df,new],ignore_index=True)
    total_df.iloc[:,4] = [i[0].text for i in total_df.iloc[:,4].values]
    total_df
    document = Document()
    for i in range(len(total_df)):
        document.add_heading(total_df.iloc[i,2],level=0)
        document.add_paragraph('Title：'+total_df.iloc[i,2], style='List Bullet')
        document.add_paragraph('Author：'+total_df.iloc[i,1], style='List Bullet')
        document.add_paragraph('Category：'+total_df.iloc[i,0], style='List Bullet')
        document.add_paragraph('Published Date：'+str(total_df.iloc[i,3]), style='List Bullet')
        document.add_paragraph('Keyword：'+total_df.iloc[i,5], style='List Bullet')
        document.add_paragraph('')
        document.add_paragraph('Abstract：\n'+total_df.iloc[i,4], style='List Bullet')
        document.add_page_break()
    document.save('demo.docx')
        

            
        
