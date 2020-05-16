from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import urllib
import tqdm
import time
import pandas as pd
from docx import Document
from docx.shared import Inches

options = webdriver.ChromeOptions()
options.add_argument('--ignore-certificate-errors')
options.add_argument('--ignore-ssl-errors')
head_url = "https://onlinelibrary.wiley.com"
if __name__ == '__main__':
    total_df = pd.DataFrame(columns=["類別","作者","標題","出版日","摘要","關鍵字"])
    driver = webdriver.Chrome(executable_path='C:\\Users\\chiaming\\Desktop\\vs\\武漢肺炎專區\\chromedriver.exe',chrome_options=options)
    for i in range(1,7,1):
        if i == 1:
            driver.get("https://onlinelibrary.wiley.com/action/doSearch?AfterYear=2020&BeforeYear=2020&SeriesKey=10970258&content=articlesChapters&countTerms=true&sortBy=Earliest&target=default")
        else:
            WebDriverWait(driver, 10).until(lambda d: d.find_element_by_css_selector(".pagination__list li:nth-child("+str(i)+") a")).click()
        target = BeautifulSoup(WebDriverWait(driver, 10).until(lambda d: d.page_source),'lxml')
        dtype = [ i.text for i in target.select(".meta__type")]
        author = []
        title = [ i.text for i in target.select(".visitable")]
        url = [ i['href'] for i in target.select(".visitable")]
        date = [ str(i)[str(i).find('</span>')+7:str(i).find('2020')+4] for i in target.select(".meta__epubDate")]
        main = []
        keyword = []
        print(i,len(dtype))
        
        for au in target.select('.comma'):
            text_author = ""
            for th in au.select('span > a > span'):
                abc= th.text.strip().replace('\n','')
                text_author = text_author + "、" + abc
            author.append(text_author[1:])

        for i in url:
            driver.get(head_url+i)
            time.sleep(2)
            abstract = BeautifulSoup(WebDriverWait(driver, 10).until(lambda d: d.page_source),'lxml')
            content = [ i.text for i in abstract.select('#section-1-en p')]
            try:
                main.append(content[0])
            except:
                main.append("無資料")
            keyword.append([ i.text for i in abstract.select(".badge-type")])
        new_keyword = []
        for b in keyword:
            text = ""
            for a in b:
                text = text +"、"+a
            print(text.strip())
            new_keyword.append(text.strip())
        driver.get("https://onlinelibrary.wiley.com/action/doSearch?AfterYear=2020&BeforeYear=2020&SeriesKey=10970258&content=articlesChapters&countTerms=true&sortBy=Earliest&target=default")
        new = {"類別":dtype,"作者":author,"標題":title,"出版日":date,"摘要":main,"關鍵字":new_keyword}
        new = pd.DataFrame(new)
        print(new.head(3))
        total_df = pd.concat([total_df,new],ignore_index=True)
    total_df.to_csv('outcome.csv',encoding='utf_8_sig')
    document = Document()
    for i in range(len(total_df)):
        if total_df.iloc[i,0] not in ['RESEARCH ARTICLE','TUTORIAL IN BIOSTATISTICS']:
            continue
        else:
            document.add_heading(total_df.iloc[i,1],level=0)
            document.add_paragraph('Title：'+total_df.iloc[i,2], style='List Bullet')
            document.add_paragraph('Author：'+total_df.iloc[i,1], style='List Bullet')
            document.add_paragraph('Category：'+total_df.iloc[i,0], style='List Bullet')
            document.add_paragraph('Published Date：'+str(total_df.iloc[i,3])[:10], style='List Bullet')
            document.add_paragraph('Keyword：'+total_df.iloc[i,5][1:], style='List Bullet')
            document.add_paragraph('')
            document.add_paragraph('Abstract：\n'+total_df.iloc[i,4], style='List Bullet')
            document.add_page_break()
    document.save('demo.docx')
